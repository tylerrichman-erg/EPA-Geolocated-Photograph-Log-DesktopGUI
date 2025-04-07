import docx
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns
import os

#Part of Page Number Methods
def create_element(name):
    return OxmlElement(name)

#Part of Page Number Methods
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)
    
#Main Page Number Method (Call This)
def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_report(
    df,
    filename_field,
    output_file_path,
    photographer,
    facility,
    inspection_date,
    image_files,
    overview_title,
    overview_text,
    overview_img_width_in,
    overview_img_folder_path,
    individual_photo_img_width_in,
    individual_imagery_folder_path,
    individual_imagery_img_width_in,
    individual_terrain_folder_path,
    individual_terrain_img_width_in,
    individual_header_end_text,
    individual_footer_beginning_text
    ):
    
    #Establish Document and open section 0 (only section)
    document = Document()
    section = document.sections[0]
    
    #Title Page
    paragraphtitle = document.add_paragraph()
    paragraphtitle.alignment = 1
    titleRun1 = paragraphtitle.add_run("Photograph Log")
    titleRun1.bold = True
    titleRun1.add_break()
    titleRun2 = paragraphtitle.add_run(facility)
    titleRun2.bold = True
    titleRun2.add_break()
    paragraphtitle.add_run("by " + photographer)
    document.add_page_break()
    
    #Overview Map Page
    mapTitle = document.add_paragraph()
    mapTitle.alignment = 1
    runMapTitle = mapTitle.add_run(overview_title)
    runMapTitle.bold = True
    runMapImage = mapTitle.add_run()
    runMapImage.add_picture(os.path.join(overview_img_folder_path, "overview.jpg"), width = Inches(individual_photo_img_width_in))
    runMapNote = mapTitle.add_run(overview_text)
    document.add_page_break()
    
    #Individual Image Pages
    photoNumber = 0
    for index, row in df.iterrows():
        photoNumber += 1
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_files[photoNumber - 1], width=Inches(individual_photo_img_width_in))
        run2 = paragraph.add_run()
        run2.add_picture(os.path.join(individual_imagery_folder_path, os.path.basename(image_files[photoNumber - 1])), width=Inches(individual_imagery_img_width_in))
        run3 = paragraph.add_run()
        run3.add_picture(os.path.join(individual_terrain_folder_path, os.path.basename(image_files[photoNumber - 1])), width=Inches(individual_terrain_img_width_in))
        paragraph.add_run("Photograph "+str(photoNumber)+"("+row[filename_field]+"):")
        document.add_page_break()
        
    #Header and Footer, Page Numbers
    header =section.header
    headertext = header.paragraphs[0]
    headertext.text = facility + " " + individual_header_end_text
    footer = section.footer
    footertext = footer.paragraphs[0]
    facilityDate = inspection_date
    footertext.text = individual_footer_beginning_text + " " + facilityDate
    footer.add_paragraph()
    add_page_number(document.sections[0].footer.paragraphs[1].add_run())
    
    #Save Output
    document.save(output_file_path)